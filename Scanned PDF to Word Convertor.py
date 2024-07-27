import os
import pytesseract
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pikepdf import Pdf
from datetime import datetime
from docx.oxml import OxmlElement

def extract_text_and_images_from_pdf(pdf_path):
    images = convert_from_path(pdf_path)
    text = ""

    for img in images:
        text += pytesseract.image_to_string(img) + "\n\n"
    
    return text, images

def create_word_doc(text, images, output_path):
    doc = Document()

    for img in images:
        img_path = "temp_image.png"
        img.save(img_path)
        doc.add_picture(img_path, width=Inches(6))
        os.remove(img_path)

    # Adding text after images to retain formatting as close as possible
    for paragraph in text.split('\n\n'):
        p = doc.add_paragraph(paragraph)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    doc.save(output_path)

def parse_pdf_datetime(pdf_datetime):
    if isinstance(pdf_datetime, str):
        try:
            return datetime.strptime(pdf_datetime, "D:%Y%m%d%H%M%S")
        except ValueError:
            return None
    return None

def retain_pdf_metadata(pdf_path, docx_path):
    with Pdf.open(pdf_path) as pdf:
        metadata = pdf.docinfo

    doc = Document(docx_path)
    core_properties = doc.core_properties

    if '/Title' in metadata:
        core_properties.title = metadata.get('/Title')
    if '/Author' in metadata:
        core_properties.author = metadata.get('/Author')
    if '/Subject' in metadata:
        core_properties.subject = metadata.get('/Subject')
    if '/Keywords' in metadata:
        core_properties.keywords = metadata.get('/Keywords')
    if '/Creator' in metadata:
        core_properties.comments = metadata.get('/Creator')
    if '/CreationDate' in metadata:
        creation_date = parse_pdf_datetime(metadata.get('/CreationDate'))
        if creation_date:
            core_properties.created = creation_date
    if '/ModDate' in metadata:
        mod_date = parse_pdf_datetime(metadata.get('/ModDate'))
        if mod_date:
            core_properties.modified = mod_date

    doc.save(docx_path)

def convert_scanned_pdf_to_word(pdf_path, output_path):
    text, images = extract_text_and_images_from_pdf(pdf_path)
    create_word_doc(text, images, output_path)
    retain_pdf_metadata(pdf_path, output_path)

# Example usage
input_pdf_path = "VTU Circular.pdf"
output_word_path = "output_docx.docx"

convert_scanned_pdf_to_word(input_pdf_path, output_word_path)
