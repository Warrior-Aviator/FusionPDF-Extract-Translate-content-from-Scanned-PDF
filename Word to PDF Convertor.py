import win32com.client
import pythoncom
import pywintypes

import os
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
import win32com.client

def convert_word_to_pdf(input_path, output_path):
    word = win32com.client.Dispatch('Word.Application')
    doc = word.Documents.Open(input_path)
    doc.SaveAs(output_path, FileFormat=17)
    doc.Close()
    word.Quit()

def extract_images_from_word(doc_path):
    doc = Document(doc_path)
    images = []

    for i, shape in enumerate(doc.inline_shapes):
        if shape.type == 3:  # Type 3 is a picture
            img = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
            image = doc.part.related_parts[img]
            image_path = f"image_{i}.png"
            with open(image_path, "wb") as f:
                f.write(image.blob)
            images.append(image_path)
    
    return images

def create_pdf_with_text_and_images(doc_path, pdf_path):
    doc = Document(doc_path)
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    text_obj = c.beginText(40, height - 40)
    
    for para in doc.paragraphs:
        for run in para.runs:
            if text_obj.getY() < 40:
                c.drawText(text_obj)
                c.showPage()
                text_obj = c.beginText(40, height - 40)
            text_obj.textLine(run.text)
    
    c.drawText(text_obj)
    
    images = extract_images_from_word(doc_path)
    for img_path in images:
        c.showPage()
        img = ImageReader(img_path)
        c.drawImage(img, 40, height - 300, width=200, height=150)
        os.remove(img_path)  # Clean up the temporary image file
    
    c.save()

# Example usage
input_word_path = "D:\College\Sem-6\Mini Project\VTU Circular.docx"
output_pdf_path = "D:\College\Sem-6\Mini Project\output_doc.pdf"

convert_word_to_pdf(input_word_path, output_pdf_path)
